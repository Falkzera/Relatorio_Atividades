# === IMPORTS ===
import os
import sys
import io
import pickle
import streamlit as st
from docx import Document as DocxDocument
from googleapiclient.http import MediaIoBaseDownload
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_chroma import Chroma
from chromadb import Client
from chromadb.config import Settings

from get_embedding_function import get_embedding_function
from Scripts.google_drive_utils import (
    authenticate_service_account,
    list_files_in_folder,
    carregar_cache_docx_do_drive,
    salvar_cache_docx_no_drive,
    enviar_arquivo_para_drive,
    baixar_arquivo_do_drive
)

# === CONSTANTES ===
ROOT_FOLDER_ID = "1z6MYXz1sy8yMkLfg6WMsIZD1nycII0rU"
ID_PASTA_CACHE = "1d0KqEyocTO1lbnWS1u7hooSVJgv5Fz6Q"
CACHE_FILENAME = "vetores_cache.pkl"

# === FUNÇÕES AUXILIARES ===

def baixar_docx(service, file_id):
    try:
        request = service.files().export_media(
            fileId=file_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        fh = io.BytesIO()
        MediaIoBaseDownload(fh, request).next_chunk()
        fh.seek(0)
        return "\n".join([p.text for p in DocxDocument(fh).paragraphs])
    except:
        try:
            request = service.files().get_media(fileId=file_id)
            fh = io.BytesIO()
            MediaIoBaseDownload(fh, request).next_chunk()
            fh.seek(0)
            return "\n".join([p.text for p in DocxDocument(fh).paragraphs])
        except:
            return ""

def coletar_arquivos_docx_do_drive():
    st.sidebar.info("🔄 Conectando ao Google Drive...")
    service = authenticate_service_account()
    cache_docx, cache_file_id = carregar_cache_docx_do_drive(service)
    novos_docs = []

    for pasta_ano in list_files_in_folder(service, ROOT_FOLDER_ID):
        for pasta_mes in list_files_in_folder(service, pasta_ano["id"]):
            for arq in list_files_in_folder(service, pasta_mes["id"]):
                nome = arq["name"]
                arq_id = arq["id"]

                if not nome.lower().endswith(".docx"):
                    continue

                if arq_id in cache_docx:
                    novos_docs.append(cache_docx[arq_id])
                else:
                    texto = baixar_docx(service, arq_id)
                    if not texto.strip():
                        continue

                    doc_info = {
                        "nome": nome,
                        "ano": pasta_ano["name"],
                        "mes": pasta_mes["name"],
                        "texto": texto,
                        "id": arq_id
                    }

                    cache_docx[arq_id] = doc_info
                    novos_docs.append(doc_info)
                    print(f"✅ Arquivo **{nome}** (ID: {arq_id}) baixado e adicionado ao cache.")

    salvar_cache_docx_no_drive(service, cache_docx, cache_file_id)
    st.success(f"✅ Total de documentos carregados: {len(novos_docs)}")
    return novos_docs

def transformar_em_documentos(dados_raw):
    return [
        Document(
            page_content=dado["texto"],
            metadata={
                "source": dado["nome"],
                "ano": dado["ano"],
                "mes": dado["mes"],
                "id_drive": dado["id"]
            }
        ) for dado in dados_raw
    ]

def dividir_em_chunks(docs):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_documents(docs)

# def salvar_chunks_no_drive(chunks):
#     local_path = "/tmp/vetores_cache.pkl"
#     with open(local_path, "wb") as f:
#         pickle.dump(chunks, f)

def salvar_chunks_no_drive(chunks_dict):
    local_path = "/tmp/vetores_cache.pkl"
    with open(local_path, "wb") as f:
        pickle.dump(chunks_dict, f)

    service = authenticate_service_account()
    enviar_arquivo_para_drive(service, local_path, CACHE_FILENAME, ID_PASTA_CACHE)

# === INDEXAÇÃO COMPLETA ===

def indexar_atas_do_drive():
    st.info("🚀 Iniciando indexação e salvamento de vetores...")
    dados_docx = coletar_arquivos_docx_do_drive()
    documentos = transformar_em_documentos(dados_docx)
    chunks = dividir_em_chunks(documentos)
    salvar_chunks_no_drive(chunks)
    st.success("✅ Cache vetorial salvo no Google Drive com sucesso!")

# === BUSCA: carregar Chroma da memória sem erro ===

@st.cache_resource()
def carregar_chroma_memoria_do_cache():
    # st.sidebar.info("📦 Carregando cache vetorial do Drive...")

    service = authenticate_service_account()
    local_path = "/tmp/vetores_cache.pkl"
    sucesso = baixar_arquivo_do_drive(service, CACHE_FILENAME, local_path, ID_PASTA_CACHE)

    if not sucesso:
        st.error("❌ Não foi possível carregar o cache vetorial do Drive.")
        return None

    try:
        with open(local_path, "rb") as f:
            chunks = pickle.load(f)

        embedding = get_embedding_function()

        import tempfile

        persist_dir = tempfile.mkdtemp()

        return Chroma.from_documents(
            documents=chunks,
            embedding=embedding,
            collection_name="atas_memoria",
            persist_directory=persist_dir
        )


    except Exception as e:
        st.error(f"❌ Erro ao carregar o cache: {e}")
        return None
